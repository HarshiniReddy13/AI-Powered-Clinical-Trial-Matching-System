"""
backend.py — Clinical Trial Matching backend
Rebuilt from trial_matching_ver5_with_pickle.ipynb (canonical source).
"""

import os, re, json
import numpy as np
from sklearn.tree import _tree
from pathlib import Path
from models import get_models as _get_models

# ---------------------------------------------------------------------------
# Lazy model globals — populated on first call to run_pipeline()
# ---------------------------------------------------------------------------
XGB              = None
DT               = None
SCALER_FULL      = None
ALL_FEAT_NAMES   = None
SELECTED_LABS_SET   = set()
SELECTED_VITALS_SET = set()

def _ensure_models_loaded():
    global XGB, DT, SCALER_FULL, ALL_FEAT_NAMES
    global SELECTED_LABS_SET, SELECTED_VITALS_SET
    if XGB is None:
        m = _get_models()
        XGB              = m["xgb"]
        DT               = m["dt"]
        SCALER_FULL      = m["scaler"]
        ALL_FEAT_NAMES   = m["feature_names"]
        SELECTED_LABS_SET   = m["selected_labs"]
        SELECTED_VITALS_SET = m["selected_vitals"]


# ---------------------------------------------------------------------------
# Trial parsing  (mirrors read_trial() in the notebook)
# ---------------------------------------------------------------------------
USABLE_STATUSES = {"RECRUITING", "NOT_YET_RECRUITING", "ENROLLING_BY_INVITATION"}

STATUS_LABEL = {
    "RECRUITING":              "Recruiting",
    "NOT_YET_RECRUITING":      "Opening Soon",
    "ENROLLING_BY_INVITATION": "By Invitation",
    "COMPLETED":               "Completed",
    "TERMINATED":              "Terminated",
    "WITHDRAWN":               "Withdrawn",
    "ACTIVE_NOT_RECRUITING":   "Active (Closed)",
    "SUSPENDED":               "Suspended",
    "UNKNOWN":                 "Unknown",
}

def _read_trial(t):
    cr  = t.get("clinical_requirements", {})
    dem = cr.get("demographics", {})
    age = dem.get("age", {})
    raw_status = t.get("overall_status", "").upper().replace(" ", "_")
    return {
        "trial_id"    : t.get("trial_id", ""),
        "title"       : t.get("title", ""),
        "conditions"  : [c.lower().strip() for c in t.get("conditions", [])],
        "phase"       : t.get("phase", []),
        "status"      : raw_status,
        "status_label": STATUS_LABEL.get(raw_status, raw_status),
        "min_age"     : age.get("min"),
        "max_age"     : age.get("max"),
        "sex"         : dem.get("sex", "ALL"),
        "lab_tests"   : cr.get("lab_tests", []),
        "vitals"      : cr.get("vitals", []),
        "scores"      : cr.get("scores", []),
        "treatments"  : cr.get("treatment_history", []),
        "inclusion"   : t.get("eligibility", {}).get("inclusion", []),
        "exclusion"   : t.get("eligibility", {}).get("exclusion", []),
    }


def load_trials(path="trials_extracted.json"):
    """Load and parse recruiting trials from JSON."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(
            f"trials_extracted.json not found at '{path}'. "
            "Make sure it is in the same directory as backend.py."
        )
    with open(p, "r", encoding="utf-8") as f:
        raw = json.load(f)
    if isinstance(raw, dict):
        raw = next(v for v in raw.values() if isinstance(v, list))
    trials = [_read_trial(t) for t in raw
              if _read_trial(t)["status"] in USABLE_STATUSES]
    if not trials:
        raise ValueError(
            "No recruiting trials found in trials_extracted.json. "
            "Check that the file contains trials with RECRUITING / "
            "NOT_YET_RECRUITING / ENROLLING_BY_INVITATION status."
        )
    return trials


# ---------------------------------------------------------------------------
# Disease hierarchy  (verbatim from notebook)
# ---------------------------------------------------------------------------
#Disease hierarchy
import re

# Disease hierarchy
DISEASE_HIERARCHY = {

    # ── ONCOLOGY ──────────────────────────────────────────────────────────────
    # Lung
    "non-small cell lung cancer" : {"parents": {"lung cancer", "lung neoplasm", "thoracic cancer"},
                                     "siblings": {"small cell lung cancer", "sclc", "large cell carcinoma"}},
    "nsclc"                      : {"parents": {"lung cancer", "lung neoplasm", "thoracic cancer"},
                                     "siblings": {"small cell lung cancer", "sclc", "large cell carcinoma"}},
    "small cell lung cancer"     : {"parents": {"lung cancer", "lung neoplasm", "thoracic cancer"},
                                     "siblings": {"non-small cell lung cancer", "nsclc", "adenocarcinoma of the lung"}},
    "sclc"                       : {"parents": {"lung cancer", "lung neoplasm"},
                                     "siblings": {"non-small cell lung cancer", "nsclc"}},
    "lung adenocarcinoma"        : {"parents": {"non-small cell lung cancer", "lung cancer"},
                                     "siblings": {"small cell lung cancer", "sclc", "squamous cell carcinoma of the lung"}},
    # Breast
    "her2-positive breast cancer"    : {"parents": {"breast cancer"},
                                         "siblings": {"triple negative breast cancer", "er-positive breast cancer", "hr-positive breast cancer"}},
    "triple negative breast cancer"  : {"parents": {"breast cancer"},
                                         "siblings": {"her2-positive breast cancer", "er-positive breast cancer", "hr-positive breast cancer"}},
    "tnbc"                           : {"parents": {"breast cancer"},
                                         "siblings": {"her2-positive breast cancer", "er-positive breast cancer"}},
    "er-positive breast cancer"      : {"parents": {"breast cancer"},
                                         "siblings": {"triple negative breast cancer", "her2-positive breast cancer"}},
    # Gynecologic
    "ovarian cancer"             : {"parents": {"gynecologic cancer", "gynecological cancer"},
                                     "siblings": {"cervical cancer", "uterine cancer", "endometrial cancer", "vulvar cancer"}},
    "recurrent ovarian cancer"   : {"parents": {"ovarian cancer", "gynecologic cancer"},
                                     "siblings": {"cervical cancer", "uterine cancer", "endometrial cancer"}},
    "cervical cancer"            : {"parents": {"gynecologic cancer"},
                                     "siblings": {"ovarian cancer", "uterine cancer", "endometrial cancer"}},
    "endometrial cancer"         : {"parents": {"gynecologic cancer", "uterine cancer"},
                                     "siblings": {"ovarian cancer", "cervical cancer"}},
    # Colorectal & GI
    "colorectal cancer"          : {"parents": {"gastrointestinal cancer", "gi cancer"},
                                     "siblings": {"gastric cancer", "pancreatic cancer", "hepatocellular carcinoma"}},
    "colon cancer"               : {"parents": {"colorectal cancer", "gastrointestinal cancer"},
                                     "siblings": {"rectal cancer", "gastric cancer", "pancreatic cancer"}},
    "rectal cancer"              : {"parents": {"colorectal cancer", "gastrointestinal cancer"},
                                     "siblings": {"colon cancer", "gastric cancer"}},
    "gastric cancer"             : {"parents": {"gastrointestinal cancer"},
                                     "siblings": {"colorectal cancer", "pancreatic cancer", "esophageal cancer"}},
    "pancreatic cancer"          : {"parents": {"gastrointestinal cancer"},
                                     "siblings": {"colorectal cancer", "gastric cancer", "hepatocellular carcinoma"}},
    "hepatocellular carcinoma"   : {"parents": {"liver cancer", "gastrointestinal cancer"},
                                     "siblings": {"cholangiocarcinoma", "colorectal cancer"}},
    # Bladder & Urologic
    "bladder cancer"             : {"parents": {"urologic cancer", "urothelial cancer"},
                                     "siblings": {"gall bladder cancer", "gallbladder cancer", "renal cell carcinoma", "prostate cancer"}},
    "gall bladder cancer"        : {"parents": {"hepatobiliary cancer"},
                                     "siblings": {"bladder cancer", "cholangiocarcinoma", "liver cancer"}},
    "gallbladder cancer"         : {"parents": {"hepatobiliary cancer"},
                                     "siblings": {"bladder cancer", "cholangiocarcinoma"}},
    "renal cell carcinoma"       : {"parents": {"kidney cancer", "urologic cancer"},
                                     "siblings": {"bladder cancer", "prostate cancer"}},
    "prostate cancer"            : {"parents": {"urologic cancer"},
                                     "siblings": {"bladder cancer", "renal cell carcinoma"}},
    # Leukemia / Lymphoma
    "acute myeloid leukemia"     : {"parents": {"leukemia", "acute leukemia", "hematologic malignancy"},
                                     "siblings": {"acute lymphoblastic leukemia", "chronic myeloid leukemia", "chronic lymphocytic leukemia"}},
    "aml"                        : {"parents": {"leukemia", "acute leukemia"},
                                     "siblings": {"all", "cml", "cll"}},
    "acute lymphoblastic leukemia": {"parents": {"leukemia", "acute leukemia", "hematologic malignancy"},
                                      "siblings": {"acute myeloid leukemia", "chronic myeloid leukemia"}},
    "all"                        : {"parents": {"leukemia", "acute leukemia"},
                                     "siblings": {"aml", "cml", "cll"}},
    "diffuse large b-cell lymphoma": {"parents": {"lymphoma", "non-hodgkin lymphoma", "hematologic malignancy"},
                                       "siblings": {"follicular lymphoma", "hodgkin lymphoma", "mantle cell lymphoma"}},
    "dlbcl"                      : {"parents": {"lymphoma", "non-hodgkin lymphoma"},
                                     "siblings": {"follicular lymphoma", "hodgkin lymphoma"}},
    # Brain
    "glioblastoma"               : {"parents": {"brain cancer", "brain tumor", "glioma"},
                                     "siblings": {"astrocytoma", "oligodendroglioma", "meningioma"}},
    "glioblastoma multiforme"    : {"parents": {"brain cancer", "glioma"},
                                     "siblings": {"astrocytoma", "oligodendroglioma"}},
    # Skin
    "melanoma"                   : {"parents": {"skin cancer"},
                                     "siblings": {"basal cell carcinoma", "squamous cell carcinoma of skin", "merkel cell carcinoma"}},
    # Thyroid
    "papillary thyroid cancer"   : {"parents": {"thyroid cancer"},
                                     "siblings": {"follicular thyroid cancer", "medullary thyroid cancer", "anaplastic thyroid cancer"}},
    "medullary thyroid cancer"   : {"parents": {"thyroid cancer"},
                                     "siblings": {"papillary thyroid cancer", "follicular thyroid cancer"}},

    # ── CARDIOVASCULAR ────────────────────────────────────────────────────────
    "heart failure"                      : {"parents": {"cardiovascular disease", "cardiac disease"},
                                             "siblings": {"coronary artery disease", "atrial fibrillation", "hypertrophic cardiomyopathy"}},
    "heart failure with reduced ejection fraction": {"parents": {"heart failure", "cardiovascular disease"},
                                                      "siblings": {"heart failure with preserved ejection fraction"}},
    "hfref"                              : {"parents": {"heart failure", "cardiovascular disease"},
                                             "siblings": {"hfpef"}},
    "heart failure with preserved ejection fraction": {"parents": {"heart failure", "cardiovascular disease"},
                                                        "siblings": {"heart failure with reduced ejection fraction"}},
    "hfpef"                              : {"parents": {"heart failure", "cardiovascular disease"},
                                             "siblings": {"hfref"}},
    "coronary artery disease"            : {"parents": {"cardiovascular disease", "cardiac disease"},
                                             "siblings": {"heart failure", "atrial fibrillation", "peripheral artery disease"}},
    "acute myocardial infarction"        : {"parents": {"coronary artery disease", "cardiovascular disease"},
                                             "siblings": {"unstable angina", "stable angina"}},
    "atrial fibrillation"                : {"parents": {"cardiac arrhythmia", "cardiovascular disease"},
                                             "siblings": {"atrial flutter", "ventricular tachycardia", "heart failure"}},
    "atrial flutter"                     : {"parents": {"cardiac arrhythmia", "cardiovascular disease"},
                                             "siblings": {"atrial fibrillation", "ventricular tachycardia"}},
    "hypertrophic cardiomyopathy"        : {"parents": {"cardiomyopathy", "cardiovascular disease"},
                                             "siblings": {"dilated cardiomyopathy", "restrictive cardiomyopathy"}},
    "dilated cardiomyopathy"             : {"parents": {"cardiomyopathy", "cardiovascular disease"},
                                             "siblings": {"hypertrophic cardiomyopathy", "restrictive cardiomyopathy"}},
    "peripheral artery disease"          : {"parents": {"cardiovascular disease", "vascular disease"},
                                             "siblings": {"coronary artery disease", "stroke"}},

    # ── NEUROLOGY ─────────────────────────────────────────────────────────────
    "alzheimer's disease"                : {"parents": {"dementia", "neurodegenerative disease", "neurological disease"},
                                             "siblings": {"vascular dementia", "lewy body dementia", "frontotemporal dementia", "parkinson's disease"}},
    "parkinson's disease"                : {"parents": {"neurodegenerative disease", "movement disorder", "neurological disease"},
                                             "siblings": {"multiple system atrophy", "progressive supranuclear palsy", "alzheimer's disease"}},
    "multiple sclerosis"                 : {"parents": {"demyelinating disease", "neurological disease", "autoimmune disease"},
                                             "siblings": {"neuromyelitis optica", "transverse myelitis"}},
    "relapsing-remitting multiple sclerosis": {"parents": {"multiple sclerosis", "demyelinating disease"},
                                                "siblings": {"primary progressive multiple sclerosis", "secondary progressive multiple sclerosis"}},
    "rrms"                               : {"parents": {"multiple sclerosis", "demyelinating disease"},
                                             "siblings": {"ppms", "spms"}},
    "epilepsy"                           : {"parents": {"neurological disease", "seizure disorder"},
                                             "siblings": {"focal epilepsy", "generalized epilepsy"}},
    "focal epilepsy"                     : {"parents": {"epilepsy", "seizure disorder"},
                                             "siblings": {"generalized epilepsy"}},
    "ischemic stroke"                    : {"parents": {"stroke", "cerebrovascular disease", "neurological disease"},
                                             "siblings": {"hemorrhagic stroke", "transient ischemic attack"}},
    "amyotrophic lateral sclerosis"      : {"parents": {"neurodegenerative disease", "motor neuron disease", "neurological disease"},
                                             "siblings": {"spinal muscular atrophy", "primary lateral sclerosis"}},
    "als"                                : {"parents": {"neurodegenerative disease", "motor neuron disease"},
                                             "siblings": {"sma", "primary lateral sclerosis"}},
    "migraine"                           : {"parents": {"headache disorder", "neurological disease"},
                                             "siblings": {"cluster headache", "tension headache"}},
    "chronic migraine"                   : {"parents": {"migraine", "headache disorder"},
                                             "siblings": {"episodic migraine", "cluster headache"}},

    # ── AUTOIMMUNE / RHEUMATOLOGIC ────────────────────────────────────────────
    "rheumatoid arthritis"               : {"parents": {"autoimmune disease", "inflammatory arthritis", "rheumatic disease"},
                                             "siblings": {"psoriatic arthritis", "ankylosing spondylitis", "systemic lupus erythematosus"}},
    "psoriatic arthritis"                : {"parents": {"autoimmune disease", "inflammatory arthritis", "rheumatic disease"},
                                             "siblings": {"rheumatoid arthritis", "ankylosing spondylitis"}},
    "ankylosing spondylitis"             : {"parents": {"autoimmune disease", "inflammatory arthritis", "spondyloarthropathy"},
                                             "siblings": {"psoriatic arthritis", "rheumatoid arthritis"}},
    "systemic lupus erythematosus"       : {"parents": {"autoimmune disease", "connective tissue disease", "rheumatic disease"},
                                             "siblings": {"rheumatoid arthritis", "sjogren's syndrome", "systemic sclerosis"}},
    "sle"                                : {"parents": {"autoimmune disease", "connective tissue disease"},
                                             "siblings": {"rheumatoid arthritis", "sjogren's syndrome"}},
    "lupus nephritis"                    : {"parents": {"systemic lupus erythematosus", "nephritis", "autoimmune disease"},
                                             "siblings": {"iga nephropathy", "focal segmental glomerulosclerosis"}},
    "psoriasis"                          : {"parents": {"autoimmune disease", "inflammatory skin disease"},
                                             "siblings": {"atopic dermatitis", "hidradenitis suppurativa"}},
    "plaque psoriasis"                   : {"parents": {"psoriasis", "autoimmune disease"},
                                             "siblings": {"pustular psoriasis", "guttate psoriasis", "psoriatic arthritis"}},
    "atopic dermatitis"                  : {"parents": {"autoimmune disease", "inflammatory skin disease", "allergic disease"},
                                             "siblings": {"psoriasis", "contact dermatitis"}},
    "inflammatory bowel disease"         : {"parents": {"autoimmune disease", "gastrointestinal disease"},
                                             "siblings": {"irritable bowel syndrome"}},
    "crohn's disease"                    : {"parents": {"inflammatory bowel disease", "autoimmune disease"},
                                             "siblings": {"ulcerative colitis"}},
    "ulcerative colitis"                 : {"parents": {"inflammatory bowel disease", "autoimmune disease"},
                                             "siblings": {"crohn's disease"}},

    # ── METABOLIC / ENDOCRINE ─────────────────────────────────────────────────
    "type 1 diabetes"                    : {"parents": {"diabetes mellitus", "autoimmune disease", "metabolic disease"},
                                             "siblings": {"type 2 diabetes", "gestational diabetes", "mody"}},
    "type 2 diabetes"                    : {"parents": {"diabetes mellitus", "metabolic disease"},
                                             "siblings": {"type 1 diabetes", "gestational diabetes", "mody"}},
    "t2dm"                               : {"parents": {"diabetes mellitus", "metabolic disease"},
                                             "siblings": {"type 1 diabetes", "gestational diabetes"}},
    "non-alcoholic steatohepatitis"      : {"parents": {"liver disease", "metabolic disease", "non-alcoholic fatty liver disease"},
                                             "siblings": {"alcoholic hepatitis", "autoimmune hepatitis"}},
    "nash"                               : {"parents": {"liver disease", "metabolic disease", "nafld"},
                                             "siblings": {"alcoholic hepatitis", "autoimmune hepatitis"}},
    "obesity"                            : {"parents": {"metabolic disease"},
                                             "siblings": {"type 2 diabetes", "metabolic syndrome"}},
    "chronic kidney disease"             : {"parents": {"renal disease", "metabolic disease"},
                                             "siblings": {"acute kidney injury", "diabetic nephropathy"}},
    "ckd"                                : {"parents": {"renal disease", "metabolic disease"},
                                             "siblings": {"acute kidney injury", "diabetic nephropathy"}},
    "hypothyroidism"                     : {"parents": {"thyroid disease", "endocrine disease"},
                                             "siblings": {"hyperthyroidism", "hashimoto's thyroiditis"}},
    "hyperthyroidism"                    : {"parents": {"thyroid disease", "endocrine disease"},
                                             "siblings": {"hypothyroidism", "graves' disease"}},

    # ── RESPIRATORY (NON-ONCOLOGIC) ───────────────────────────────────────────
    "asthma"                             : {"parents": {"respiratory disease", "obstructive lung disease", "allergic disease"},
                                             "siblings": {"copd", "allergic rhinitis", "bronchiectasis"}},
    "severe asthma"                      : {"parents": {"asthma", "respiratory disease"},
                                             "siblings": {"eosinophilic asthma", "allergic asthma"}},
    "copd"                               : {"parents": {"respiratory disease", "obstructive lung disease"},
                                             "siblings": {"asthma", "bronchiectasis", "pulmonary fibrosis"}},
    "idiopathic pulmonary fibrosis"      : {"parents": {"pulmonary fibrosis", "interstitial lung disease", "respiratory disease"},
                                             "siblings": {"non-specific interstitial pneumonia", "sarcoidosis"}},
    "ipf"                                : {"parents": {"pulmonary fibrosis", "interstitial lung disease"},
                                             "siblings": {"nsip", "sarcoidosis"}},

    # ── INFECTIOUS DISEASE ────────────────────────────────────────────────────
    "hiv infection"                      : {"parents": {"infectious disease", "viral infection", "retroviral infection"},
                                             "siblings": {"hepatitis b", "hepatitis c"}},
    "aids"                               : {"parents": {"hiv infection", "infectious disease"},
                                             "siblings": {"hepatitis b", "hepatitis c"}},
    "hepatitis b"                        : {"parents": {"viral hepatitis", "liver disease", "infectious disease"},
                                             "siblings": {"hepatitis c", "hepatitis a", "hiv infection"}},
    "hepatitis c"                        : {"parents": {"viral hepatitis", "liver disease", "infectious disease"},
                                             "siblings": {"hepatitis b", "hepatitis a", "hiv infection"}},
    "tuberculosis"                       : {"parents": {"infectious disease", "bacterial infection", "mycobacterial infection"},
                                             "siblings": {"non-tuberculous mycobacterial infection", "pneumonia"}},
    "sepsis"                             : {"parents": {"infectious disease", "critical illness"},
                                             "siblings": {"septic shock", "bacteremia"}},
    "septic shock"                       : {"parents": {"sepsis", "infectious disease", "critical illness"},
                                             "siblings": {"bacteremia"}},

    # ── PSYCHIATRIC ───────────────────────────────────────────────────────────
    "major depressive disorder"          : {"parents": {"mood disorder", "psychiatric disorder"},
                                             "siblings": {"bipolar disorder", "persistent depressive disorder", "anxiety disorder"}},
    "mdd"                                : {"parents": {"mood disorder", "psychiatric disorder"},
                                             "siblings": {"bipolar disorder", "anxiety disorder"}},
    "bipolar disorder"                   : {"parents": {"mood disorder", "psychiatric disorder"},
                                             "siblings": {"major depressive disorder", "cyclothymia"}},
    "bipolar i disorder"                 : {"parents": {"bipolar disorder", "mood disorder"},
                                             "siblings": {"bipolar ii disorder", "cyclothymia"}},
    "schizophrenia"                      : {"parents": {"psychotic disorder", "psychiatric disorder"},
                                             "siblings": {"schizoaffective disorder", "delusional disorder"}},
    "generalized anxiety disorder"       : {"parents": {"anxiety disorder", "psychiatric disorder"},
                                             "siblings": {"panic disorder", "social anxiety disorder", "major depressive disorder"}},
    "post-traumatic stress disorder"     : {"parents": {"trauma-related disorder", "psychiatric disorder"},
                                             "siblings": {"acute stress disorder", "adjustment disorder"}},
    "ptsd"                               : {"parents": {"trauma-related disorder", "psychiatric disorder"},
                                             "siblings": {"acute stress disorder", "adjustment disorder"}},
}

# ── GENERIC TERMS ─────────────────────────────────────────────────────────────
# Expanded beyond oncology to cover any therapeutic area
GENERIC_DISEASE_TERMS = {
    # Oncology generics
    "cancer", "tumor", "neoplasm", "malignancy", "malignant",
    "solid tumor", "advanced cancer", "metastatic cancer",
    "solid", "advanced", "metastatic",
    # Cross-therapeutic-area generics
    "disease", "disorder", "condition", "syndrome", "infection",
    "illness", "chronic", "acute", "severe", "mild", "moderate",
    # Symptoms (not specific enough for matching)
    "pain", "nausea", "fatigue", "inflammation", "fever",
}

def _is_generic_only(condition_str):
    words = set(condition_str.lower().strip().split())
    return words.issubset(GENERIC_DISEASE_TERMS)


def _match_single_condition(pc, tc):
    """
    Score one patient condition (pc) against one trial condition (tc).
    Returns float 0.0–1.0.

    Order of checks (CRITICAL — order matters):
      1. Exact match                      → 1.0
      2. Hierarchy check FIRST            → block siblings (0.0) or allow parents (0.5)
      3. Word-boundary substring check    → prevents 'bladder' ⊂ 'gall bladder'
      4. Generic trial                    → 0.2
      5. Word-overlap fallback            → 0.0–1.0
    """
    pc = pc.lower().strip()
    tc = tc.lower().strip()

    # 1. Exact match
    if pc == tc:
        return 1.0

    # 2. Hierarchy check FIRST (prevents sibling false positives)
    if pc in DISEASE_HIERARCHY:
        hier = DISEASE_HIERARCHY[pc]
        # Block siblings — hard 0
        if (tc in hier["siblings"]
                or any(re.search(r'\b' + re.escape(s) + r'\b', tc)
                       for s in hier["siblings"])):
            return 0.0
        # Parent match — partial credit
        if (tc in hier["parents"]
                or any(re.search(r'\b' + re.escape(p) + r'\b', tc)
                       for p in hier["parents"])):
            return 0.5

    # Also check reverse hierarchy (trial condition is the subtype)
    if tc in DISEASE_HIERARCHY:
        hier = DISEASE_HIERARCHY[tc]
        if (pc in hier["siblings"]
                or any(re.search(r'\b' + re.escape(s) + r'\b', pc)
                       for s in hier["siblings"])):
            return 0.0
        if (pc in hier["parents"]
                or any(re.search(r'\b' + re.escape(p) + r'\b', pc)
                       for p in hier["parents"])):
            return 0.7   # patient is parent, trial is specific subtype — higher than reverse

    # 3. Word-boundary substring check (FIX: 'bladder' won't match 'gall bladder')
    pc_pattern = r'\b' + re.escape(pc) + r'\b'
    tc_pattern = r'\b' + re.escape(tc) + r'\b'
    if re.search(pc_pattern, tc) or re.search(tc_pattern, pc):
        return 1.0

    # 4. Generic trial → low score
    if _is_generic_only(tc):
        return 0.2

    # 5. Word-overlap fallback (stopwords expanded beyond oncology)
    STOP = {
        # Oncology
        "cancer", "tumor", "neoplasm", "malignant",
        # General medical
        "disease", "disorder", "syndrome", "condition", "infection",
        "chronic", "acute", "severe", "advanced", "recurrent",
        "primary", "secondary", "metastatic",
        # Function words
        "the", "of", "in", "and", "with",
    }
    pc_words = set(pc.split()) - STOP
    tc_words = set(tc.split()) - STOP
    if pc_words and tc_words:
        overlap = len(pc_words & tc_words) / max(len(pc_words), len(tc_words))
        if overlap >= 0.8:
            return 1.0
        if overlap >= 0.5:
            return 0.6
        if overlap >= 0.3:
            return 0.3

    return 0.0


def condition_overlap(p_conds, t_conds):
    """Average best-match score across all trial conditions."""
    if not t_conds or not p_conds:
        return 0.0
    total = 0.0
    for tc in t_conds:
        best = max((_match_single_condition(pc, tc) for pc in p_conds), default=0.0)
        total += best
    return total / len(t_conds)


# ---------------------------------------------------------------------------
# Feature engineering  (verbatim from notebook)
# ---------------------------------------------------------------------------
def _score_criterion(patient_val, op, threshold, lo, hi):
    if patient_val is None:
        return -1, 0.0
    if op == "range" and lo is not None and hi is not None:
        if lo <= patient_val <= hi:
            return 1, float(min(patient_val - lo, hi - patient_val))
        return 0, float((patient_val - lo) if patient_val < lo else (hi - patient_val))
    if threshold is None:
        return -1, 0.0
    margin = (patient_val - threshold if op in (">=", ">")
              else threshold - patient_val if op in ("<=", "<")
              else -abs(patient_val - threshold))
    return (1 if margin >= 0 else 0), float(margin)


def _get_patient_val(p_dict, trial_name):
    tn = trial_name.lower().strip()
    if tn in p_dict:
        v = p_dict[tn]; return v.get("value") if isinstance(v, dict) else v
    for pk in p_dict:
        pk_l = pk.lower().strip()
        if pk_l in tn or tn in pk_l:
            v = p_dict[pk]; return v.get("value") if isinstance(v, dict) else v
    tn_words = set(tn.replace("-", " ").split())
    for pk in p_dict:
        pk_words = set(pk.lower().replace("-", " ").split())
        if tn_words & pk_words:
            v = p_dict[pk]; return v.get("value") if isinstance(v, dict) else v
    return None


def build_features(patient, trial):
    f = {}

    # Age
    p_age  = patient.get("age")
    t_min, t_max = trial["min_age"], trial["max_age"]
    if p_age is not None:
        below = max(0.0, (t_min or 0) - p_age)
        above = max(0.0, p_age - (t_max or 999))
        f["age_ok"]        = 1.0 if (below == 0 and above == 0) else 0.0
        f["age_below_min"] = float(below)
        f["age_above_max"] = float(above)
    else:
        f["age_ok"] = f["age_below_min"] = f["age_above_max"] = -1.0
    f["has_age_limit"] = float(t_min is not None or t_max is not None)

    # Sex
    t_sex = trial["sex"]
    p_sex = patient.get("sex", "UNKNOWN")
    f["sex_ok"] = (1.0  if t_sex in ("ALL", None)
                   else -1.0 if p_sex == "UNKNOWN"
                   else  1.0 if p_sex == t_sex else 0.0)

    # Conditions
    overlap = condition_overlap(patient.get("conditions", []), trial["conditions"])
    f["condition_overlap"]  = float(overlap)
    f["n_trial_conditions"] = float(len(trial["conditions"]))
    generic_count = sum(1 for tc in trial["conditions"] if _is_generic_only(tc))
    f["trial_is_generic"] = float(generic_count / max(len(trial["conditions"]), 1))

    # ECOG
    p_ecog   = patient.get("ecog")
    ecog_req = next((s for s in trial["scores"]
                     if "ecog" in s.get("name", "").lower()
                     or "performance" in s.get("name", "").lower()), None)
    f["has_ecog_req"] = 1.0 if ecog_req else 0.0
    if ecog_req:
        met, margin = _score_criterion(
            p_ecog, ecog_req.get("operator"),
            ecog_req.get("value"), ecog_req.get("min"), ecog_req.get("max"))
        f["ecog_ok"] = float(met); f["ecog_margin"] = float(margin)
    else:
        f["ecog_ok"] = f["ecog_margin"] = -1.0

    # Labs
    p_labs = patient.get("lab_tests", {})
    lab_met, lab_margins, lab_total = [], [], 0
    for lt in trial["lab_tests"]:
        name = lt.get("test_name", "").lower().strip()
        if not name: continue
        lab_total += 1
        p_val = _get_patient_val(p_labs, name)
        met, margin = _score_criterion(
            float(p_val) if p_val is not None else None,
            lt.get("operator"), lt.get("value"), lt.get("min"), lt.get("max"))
        ref  = lt.get("value") or lt.get("max") or lt.get("min") or 1.0
        norm = float(margin / ref) if ref else 0.0
        if met != -1:
            lab_met.append(met); lab_margins.append(norm)
        if name in SELECTED_LABS_SET:
            key = name.replace(" ", "_").replace("-", "_")[:40]
            f[f"lab_{key}_ok"] = float(met); f[f"lab_{key}_margin"] = norm
    f["lab_met_ratio"]  = (sum(1 for v in lab_met if v == 1) / len(lab_met) if lab_met else -1.0)
    f["lab_avg_margin"] = (sum(lab_margins) / len(lab_margins) if lab_margins else -1.0)
    f["lab_total_req"]  = float(lab_total)

    # Vitals
    p_vitals = patient.get("vitals", {})
    vit_met, vit_margins = [], []
    for v in trial["vitals"]:
        name = v.get("name", "").lower().strip()
        if not name: continue
        p_val = _get_patient_val(p_vitals, name)
        met, margin = _score_criterion(
            float(p_val) if p_val is not None else None,
            v.get("operator"), v.get("value"), v.get("min"), v.get("max"))
        ref  = v.get("value") or v.get("max") or 1.0
        norm = float(margin / ref) if ref else 0.0
        if met != -1:
            vit_met.append(met); vit_margins.append(norm)
        if name in SELECTED_VITALS_SET:
            key = name.replace(" ", "_").replace("-", "_")[:40]
            f[f"vital_{key}_ok"] = float(met); f[f"vital_{key}_margin"] = norm
    f["vital_met_ratio"]  = (sum(1 for v in vit_met if v == 1) / len(vit_met) if vit_met else -1.0)
    f["vital_avg_margin"] = (sum(vit_margins) / len(vit_margins) if vit_margins else -1.0)

    # Prior treatments
    p_tx = [tx_item.lower() for tx_item in patient.get("prior_treatments", [])]
    violations = sum(
        1 for req in trial["treatments"]
        if not req.get("allowed", True)
        and any(req.get("type", "").lower() in pt or pt in req.get("type", "").lower()
                for pt in p_tx)
    )
    n_tx = len(trial["treatments"])
    f["tx_ok"]    = 1.0 if violations == 0 else max(0.0, 1.0 - violations / max(n_tx, 1))
    f["n_tx_req"] = float(n_tx)

    phase = trial["phase"]
    try:
        f["phase"] = float(phase[0].replace("PHASE", "").replace("_", "").strip()) if phase else 0.0
    except (ValueError, IndexError):
        f["phase"] = 0.0

    return f


# ---------------------------------------------------------------------------
# Scoring  (verbatim from notebook)
# ---------------------------------------------------------------------------
MIN_SCORE_THRESHOLD   = 0.25
MIN_CONDITION_OVERLAP = 0.50
LAB_VITAL_THRESHOLD   = 0.85
LAB_VITAL_MIN_REQ     = 3

SAFETY_DISCLAIMER = (
    "These trial matches are AI-generated suggestions and must be reviewed by "
    "a qualified medical professional before any clinical decision."
)

SUPPORTIVE_TRIAL_KEYWORDS = [
    "survey", "quality of life", "pain", "adherence", "digital",
    "artificial intelligence", "questionnaire", "observational", "registry",
    "patient reported", "fatigue", "nausea", "symptom management",
]
BROAD_ONCOLOGY_KEYWORDS = [
    "solid tumor", "advanced cancer", "any cancer", "hematologic malignancy",
    "pan-tumor", "tumor agnostic", "all cancer types",
]

def _is_supportive_trial(title):
    return any(kw in title.lower() for kw in SUPPORTIVE_TRIAL_KEYWORDS)

def _trial_tier(title, is_generic):
    if _is_supportive_trial(title): return "supportive"
    if is_generic > 0.5 or any(kw in title.lower() for kw in BROAD_ONCOLOGY_KEYWORDS):
        return "broad_oncology"
    return "treatment"


def compute_pair_score(row, trial_title=''):
    age    = row.get("age_ok",            -1.0)
    sex    = row.get("sex_ok",             1.0)
    cond   = row.get("condition_overlap",   0.0)
    tx     = row.get("tx_ok",              1.0)
    ecog_h = row.get("has_ecog_req",       0.0)
    ecog_m = row.get("ecog_ok",           -1.0)
    generic= row.get("trial_is_generic",   0.0)
    hard_fail = (age == 0.0) or (sex == 0.0)

    if cond < MIN_CONDITION_OVERLAP:
        return round(min(cond * 0.20, 0.15), 4)

    s  = 0.25 if age == 1.0  else (0.10 if age == -1.0  else 0.0)
    s += 0.15 if sex == 1.0  else (0.07 if sex == -1.0  else 0.0)
    s += 0.20 * min(float(cond), 1.0)
    s += 0.10 if tx >= 0.8   else (0.05 if tx >= 0.5    else 0.0)
    if ecog_h == 1.0:
        s += 0.10 if ecog_m == 1.0 else (0.05 if ecog_m == -1.0 else 0.0)
    else:
        s += 0.05

    lab_ratio   = row.get("lab_met_ratio",  -1.0)
    vital_ratio = row.get("vital_met_ratio", -1.0)
    lab_req     = row.get("lab_total_req",    0.0)

    if lab_req >= LAB_VITAL_MIN_REQ and lab_ratio >= 0:
        if lab_ratio < LAB_VITAL_THRESHOLD:
            s = min(s, 0.35)

    if lab_ratio  >= 0: s += 0.15 * float(lab_ratio)
    if vital_ratio >= 0: s += 0.05 * float(vital_ratio)

    s *= (1.0 - 0.4 * float(generic))
    if hard_fail:     s = min(s, 0.30)
    if cond < 0.70:   s = max(s - 0.20, 0.0)

    tier = _trial_tier(trial_title, generic)
    if tier == "supportive":       s -= 0.20
    elif tier == "broad_oncology": s -= 0.05
    return round(max(s, 0.0), 4)


# ---------------------------------------------------------------------------
# Explanation helpers  (verbatim from notebook)
# ---------------------------------------------------------------------------
def dt_path_explanation(x_row, feature_names, dt_model):
    tree_ = dt_model.tree_
    node  = 0
    rules = []
    while tree_.feature[node] != _tree.TREE_UNDEFINED:
        idx       = tree_.feature[node]
        feat      = feature_names[idx]
        thr       = tree_.threshold[node]
        val       = float(x_row[idx])
        direction = "<=" if val <= thr else ">"
        node      = (tree_.children_left[node] if val <= thr else tree_.children_right[node])
        rules.append(_rule_to_english(feat, direction, val, thr))
    leaf     = tree_.value[node][0]
    decision = "ELIGIBLE" if np.argmax(leaf) == 1 else "NOT ELIGIBLE"
    conf     = float(np.max(leaf) / np.sum(leaf))
    return rules, decision, round(conf, 3)


def _rule_to_english(feat, direction, val, thr):
    passed = direction == ">"
    parts  = feat.split("_")
    _map = {
        "age_ok":            (" Age is within trial range",           " Age is outside trial range"),
        "age_below_min":     (" Patient too young",                   " Meets minimum age"),
        "age_above_max":     (" Patient too old",                     " Meets maximum age"),
        "sex_ok":            (" Sex matches trial",                   " Sex mismatch"),
        "condition_overlap": (" Diagnosis overlaps trial conditions", " Low diagnosis overlap"),
        "trial_is_generic":  ("  Trial targets generic disease",      " Trial is disease-specific"),
        "ecog_ok":           (" ECOG score meets requirement",        " ECOG fails requirement"),
        "ecog_margin":       (" ECOG comfortably within limit",       "  ECOG near limit"),
        "tx_ok":             (" Treatment history compatible",        " Treatment history conflict"),
        "lab_met_ratio":     (" Most lab criteria met",               " Many lab criteria not met"),
        "vital_met_ratio":   (" Most vital criteria met",             " Many vital criteria not met"),
    }
    if feat in _map:
        return _map[feat][0 if passed else 1]
    if feat.startswith("lab_"):
        name = " ".join(parts[1:-1]).upper(); sfx = parts[-1]
        return (f" {name} meets criterion" if (sfx == "ok" and passed)
                else f" {name} fails criterion" if sfx == "ok"
                else f" {name} in range" if passed else f"⚠  {name} near limit")
    if feat.startswith("vital_"):
        name = " ".join(parts[1:-1]).upper(); sfx = parts[-1]
        return (f" {name} meets criterion" if (sfx == "ok" and passed)
                else f" {name} fails criterion" if sfx == "ok"
                else f" {name} in range" if passed else f"⚠  {name} near limit")
    return f"{'✅' if passed else '❌'} {feat} {'>' if passed else '<='} {thr:.3f} (val={val:.3f})"


def criteria_breakdown(patient, trial, feat_row):
    bd = {}
    age_ok = feat_row.get("age_ok", -1)
    bd["Age"] = {
        "status":  "PASS" if age_ok == 1 else "FAIL" if age_ok == 0 else "UNKNOWN",
        "patient": patient.get("age"),
        "trial":   f"{trial['min_age']} – {trial['max_age']}",
    }
    sex_ok = feat_row.get("sex_ok", -1)
    bd["Sex"] = {
        "status":  "PASS" if sex_ok == 1 else "FAIL" if sex_ok == 0 else "UNKNOWN",
        "patient": patient.get("sex"),
        "trial":   trial["sex"],
    }
    overlap = feat_row.get("condition_overlap", 0)
    generic = feat_row.get("trial_is_generic", 0)
    bd["Disease match"] = {
        "status":  ("PASS" if overlap >= 0.8 else "PARTIAL" if overlap >= MIN_CONDITION_OVERLAP else "FAIL"),
        "patient": ", ".join(patient.get("conditions", [])),
        "trial":   ", ".join(trial["conditions"][:4]),
        "overlap": f"{overlap*100:.0f}%",
        "generic": f"{generic*100:.0f}% generic",
    }
    if feat_row.get("has_ecog_req", 0) == 1:
        ecog_ok = feat_row.get("ecog_ok", -1)
        bd["ECOG"] = {
            "status":  "PASS" if ecog_ok == 1 else "FAIL" if ecog_ok == 0 else "UNKNOWN",
            "patient": patient.get("ecog"),
        }
    lab_ratio = feat_row.get("lab_met_ratio", -1)
    lab_req   = feat_row.get("lab_total_req", 0)
    if lab_ratio >= 0:
        threshold_note = (f" [≥85% gate: {'PASS' if lab_ratio >= LAB_VITAL_THRESHOLD else 'FAIL'}]"
                          if lab_req >= LAB_VITAL_MIN_REQ else "")
        bd["Lab tests"] = {
            "status":  "PASS" if lab_ratio >= 0.8 else "PARTIAL" if lab_ratio >= 0.5 else "FAIL",
            "patient": f"{lab_ratio*100:.0f}% of {int(lab_req)} required labs met{threshold_note}",
        }
    vital_ratio = feat_row.get("vital_met_ratio", -1)
    if vital_ratio >= 0:
        bd["Vitals"] = {
            "status":  "PASS" if vital_ratio >= 0.8 else "PARTIAL" if vital_ratio >= 0.5 else "FAIL",
            "patient": f"{vital_ratio*100:.0f}% of required vitals met",
        }
    tx_ok = feat_row.get("tx_ok", 1)
    bd["Treatment history"] = {
        "status":  "PASS" if tx_ok >= 0.8 else "PARTIAL" if tx_ok >= 0.5 else "FAIL",
        "patient": ", ".join(patient.get("prior_treatments", [])) or "none",
    }
    return bd


# ---------------------------------------------------------------------------
# Patient matching  (verbatim from notebook)
# ---------------------------------------------------------------------------
def match_patient(patient, trials, xgb_model, dt_model, scaler,
                  feature_names, top_n=10):
    feat_rows   = [build_features(patient, t) for t in trials]
    rule_scores = np.array([
        compute_pair_score(feat_rows[i], trials[i]["title"])
        for i in range(len(trials))
    ])

    valid_mask  = rule_scores >= MIN_SCORE_THRESHOLD
    valid_count = int(valid_mask.sum())

    if valid_count == 0:
        return {
            "patient_id":      patient["patient_id"],
            "age":             patient.get("age"),
            "sex":             patient.get("sex"),
            "conditions":      patient.get("conditions"),
            "ecog":            patient.get("ecog"),
            "total_trials":    len(trials),
            "top_matches":     [],
            "no_match":        True,
            "no_match_reason": (
                f"No recruiting trial found matching patient's health profile. "
                f"All {len(trials):,} evaluated trials scored below minimum threshold "
                f"({MIN_SCORE_THRESHOLD*100:.0f}%). Consider broadening search criteria."
            ),
            "disclaimer": SAFETY_DISCLAIMER,
        }

    ranked  = np.argsort(rule_scores)[::-1]
    top_idx = ranked[:max(top_n, 20)]

    X_top = np.array([
        [float(feat_rows[i].get(k, 0.0) or 0.0) for k in feature_names]
        for i in top_idx
    ])
    X_top_scaled = scaler.transform(X_top)
    xgb_preds    = xgb_model.predict(X_top_scaled)
    xgb_probs    = xgb_model.predict_proba(X_top_scaled)[:, 1]

    results, result_rank = [], 0
    for arr_pos, global_idx in enumerate(top_idx):
        if result_rank >= top_n: break
        trial      = trials[global_idx]
        rule_score = float(rule_scores[global_idx])
        if rule_score < MIN_SCORE_THRESHOLD: continue

        xgb_prob    = float(xgb_probs[arr_pos])
        final_score = round(0.7 * rule_score + 0.3 * xgb_prob, 4)
        score_pct   = round(final_score * 100, 1)
        result_rank += 1

        label = (
            "POTENTIAL MATCH (Requires Doctor Review)"  if final_score >= 0.75 else
            "POSSIBLE MATCH (Requires Doctor Review)"   if final_score >= 0.50 else
            "LOW MATCH (Requires Doctor Review)"        if final_score >= 0.30 else
            "UNLIKELY MATCH"
        )
        tier = _trial_tier(trial["title"], feat_rows[global_idx].get("trial_is_generic", 0.0))

        path, dt_dec, dt_conf = dt_path_explanation(X_top_scaled[arr_pos], feature_names, dt_model)
        breakdown = criteria_breakdown(patient, trial, feat_rows[global_idx])
        passes    = [k for k, v in breakdown.items() if v.get("status") == "PASS"]
        partials  = [k for k, v in breakdown.items() if v.get("status") == "PARTIAL"]
        fails     = [k for k, v in breakdown.items() if v.get("status") == "FAIL"]

        if final_score >= 0.50:
            summary = (f"Patient {patient['patient_id']} scored {score_pct}% — "
                       f"{len(passes)} criteria pass"
                       + (f", {len(partials)} partial" if partials else "")
                       + (". Issues: " + ", ".join(fails) + "." if fails
                          else ". No disqualifying criteria found."))
        else:
            summary = (f"Patient {patient['patient_id']} unlikely ({score_pct}%). "
                       + ("Failing: " + ", ".join(fails) + "." if fails else "Low disease overlap."))

        results.append({
            "rank":         result_rank,
            "trial_id":     trial["trial_id"],
            "title":        trial["title"][:100],
            "score":        score_pct,
            "rule_score":   round(rule_score * 100, 1),
            "xgb_prob":     round(xgb_prob, 4),
            "xgb_verdict":  "ELIGIBLE" if xgb_preds[arr_pos] == 1 else "NOT ELIGIBLE",
            "label":        label,
            "trial_tier":   tier,
            "phase":        trial["phase"],
            "status":       trial["status"],
            "status_label": trial["status_label"],
            "explanation": {
                "summary":            summary,
                "criteria_breakdown": breakdown,
                "decision_tree_path": path,
                "dt_verdict":         dt_dec,
                "dt_confidence":      f"{dt_conf*100:.0f}%",
                "score_breakdown": {
                    "final_score (hybrid)": f"{score_pct}%  (= 0.7 x rule + 0.3 x xgb)",
                    "rule_score":  f"{rule_score*100:.1f}%",
                    "xgb_prob":    f"{xgb_prob:.3f}",
                    "age":         str(feat_rows[global_idx].get("age_ok", -1)),
                    "sex":         str(feat_rows[global_idx].get("sex_ok", -1)),
                    "condition":   f"{feat_rows[global_idx].get('condition_overlap', 0)*100:.0f}% overlap",
                    "generic":     f"{feat_rows[global_idx].get('trial_is_generic', 0)*100:.0f}% generic",
                    "trial_tier":  tier,
                    "ecog":        str(feat_rows[global_idx].get("ecog_ok", -1)),
                    "labs":        (f"{feat_rows[global_idx].get('lab_met_ratio', 0)*100:.0f}% met"
                                    if feat_rows[global_idx].get("lab_met_ratio", -1) >= 0 else "n/a"),
                    "vitals":      (f"{feat_rows[global_idx].get('vital_met_ratio', 0)*100:.0f}% met"
                                    if feat_rows[global_idx].get("vital_met_ratio", -1) >= 0 else "n/a"),
                    "treatment":   f"{feat_rows[global_idx].get('tx_ok', 1)*100:.0f}% compatible",
                    "trial_status": trial["status_label"],
                },
            },
        })
    
    results.sort(key=lambda x: x["score"], reverse=True)
    for i, r in enumerate(results, 1):
        r["rank"] = i

    return {
        "patient_id":   patient["patient_id"],
        "age":          patient.get("age"),
        "sex":          patient.get("sex"),
        "conditions":   patient.get("conditions"),
        "ecog":         patient.get("ecog"),
        "total_trials": len(trials),
        "valid_trials": valid_count,
        "no_match":     False,
        "top_matches":  results,
        "disclaimer":   SAFETY_DISCLAIMER,
    }


# ---------------------------------------------------------------------------
# NLP — text extraction  (verbatim from notebook)
# ---------------------------------------------------------------------------
import spacy
NLP = spacy.blank("en")
NLP.add_pipe("sentencizer")

from transformers import pipeline as hf_pipeline
try:
    BERT_NER = hf_pipeline(
        "token-classification",
        model                = "dmis-lab/biobert-base-cased-v1.2",
        aggregation_strategy = "simple",
        device               = -1,
    )
except Exception:
    BERT_NER = None


def extract_text(file_path):
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":                                  return _read_pdf(file_path)
    elif ext in (".jpg", ".jpeg", ".png", ".webp", ".tiff"): return _read_image(file_path)
    elif ext == ".docx":                               return _read_docx(file_path)
    else:
        with open(file_path, encoding="utf-8", errors="ignore") as f: return f.read()

def _read_pdf(path):
    import fitz
    doc, pages = fitz.open(path), []
    for page in doc:
        text = page.get_text("text")
        pages.append(text if text.strip() else _ocr_bytes(page.get_pixmap(dpi=200).tobytes("png")))
    doc.close()
    return "\n".join(pages)

def _read_image(path):
    from PIL import Image; import pytesseract
    return pytesseract.image_to_string(Image.open(path))

def _read_docx(path):
    from docx import Document
    doc   = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells: lines.append("  |  ".join(cells))
    return "\n".join(lines)

def _ocr_bytes(img_bytes):
    try:
        from PIL import Image; import pytesseract, io
        return pytesseract.image_to_string(Image.open(io.BytesIO(img_bytes)))
    except Exception: return ""


def preprocess(text):
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = text.replace("\u2265", ">=").replace("\u2264", "<=")
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    for pat, rep in [
        (r"\bg/dl\b","g/dL"),(r"\bmg/dl\b","mg/dL"),(r"\bu/l\b","U/L"),
        (r"\biu/l\b","IU/L"),(r"\bml/min\b","mL/min"),(r"\bmmhg\b","mmHg"),
        (r"\bkg/m2\b","kg/m2"),(r"\bx\s*10\^9/l\b","x10^9/L"),
        (r"\bng/ml\b","ng/mL"),(r"\bmeq/l\b","mEq/L"),
    ]:
        text = re.sub(pat, rep, text, flags=re.I)
    return text.strip()


def _get_bert_ents(text):
    if not BERT_NER: return []
    try: return BERT_NER(text[:8000])
    except Exception: return []

def extract_patient_id(text):
    for pat in [
        r"(?:patient\s*(?:id|name|no)|mrn|pid)[:\s#]+([A-Za-z0-9\-\s]{2,30})",
        r"name[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)",
    ]:
        m = re.search(pat, text, re.I)
        if m:
            val = m.group(1).strip()
            if 2 < len(val) < 50: return val
    return "PT_EXTRACTED"

def extract_demographics(text):
    age = None
    for pat in [
        r"\bage[:\s]+(\d{1,3})\s*(?:years?|yrs?)?\b",
        r"\b(\d{1,3})[\s-]*year[\s-]*old\b",
        r"\b(\d{1,3})\s*(?:yo|y\.o\.)\b",
        r"\bpatient[,\s]+(?:a\s+)?(\d{1,3})[\s-]*year",
    ]:
        m = re.search(pat, text, re.I)
        if m:
            val = next((g for g in m.groups() if g), None)
            if val and 1 <= int(val) <= 120: age = int(val); break
    f  = len(re.findall(r"\b(female|woman|she\b|her\b)\b", text, re.I))
    ml = len(re.findall(r"\b(male|man\b|he\b|his\b)\b",   text, re.I))
    sex = "FEMALE" if f > ml else "MALE" if ml > f else None
    return age, sex

def extract_conditions(text, bert_ents):
    conditions = set()
    for ent in bert_ents:
        lbl = ent.get("entity_group", ent.get("label", "")).upper()
        if any(x in lbl for x in ("DIS", "DISEASE", "CANCER")):
            name = ent["word"].strip().lower()
            if len(name) > 3: conditions.add(name)
    for pat in [
        r"(?:primary\s+)?(?:diagnosis|impression|assessment)[:\s]+([^\n]{5,80})",
        r"(?:diagnosed with|presents with|known case of)[:\s]+([^\n,;]{5,60})",
    ]:
        for m in re.finditer(pat, text, re.I):
            val = re.sub(r"^(diagnosis|impression|assessment)[:\s]+",
                         "", m.group(1).strip().lower(), flags=re.I).strip()
            if 3 < len(val) < 80: conditions.add(val)
    return sorted(conditions)

def extract_ecog(text):
    for pat in [
        r"ecog\s*(?:performance\s*status|ps|score)?[:\s=]+(\d)",
        r"performance\s*status[:\s=]+(\d)",
        r"\bps\s*[:\s=]+(\d)(?:\s*/\s*\d)?",
    ]:
        m = re.search(pat, text, re.I)
        if m and 0 <= int(m.group(1)) <= 4: return int(m.group(1))
    m = re.search(r"karnofsky[:\s=]+(\d{2,3})", text, re.I)
    if m:
        k = int(m.group(1))
        return 0 if k>=90 else 1 if k>=70 else 2 if k>=50 else 3 if k>=30 else 4
    return None

_UNIT_RE = re.compile(
    r"([\d]+(?:\.\d+)?)\s*"
    r"(g/dL|g/L|mg/dL|mg/L|mmol/L|µmol/L|nmol/L|"
    r"x10\^9/L|x10\^3/µL|/mm3|U/L|IU/L|IU/mL|"
    r"ng/mL|pg/mL|µg/mL|mEq/L|mL/min(?:/1\.73\s*m2)?|"
    r"copies/mL|cells/mm3|ng/dL|pmol/L|%)", re.I,
)
_SKIP = {"page","date","time","ref","result","normal","flag","unit","range",
         "low","high","critical","reference","method","test","value",
         "comment","collected","specimen","ordered"}

def extract_lab_tests(text, bert_ents):
    labs = {}
    ner_names = set()
    for ent in bert_ents:
        lbl = ent.get("entity_group", ent.get("label", "")).upper()
        if any(x in lbl for x in ("CHEM", "LAB", "CHEMICAL")):
            ner_names.add(ent["word"].strip().lower())
    for line in text.split("\n"):
        line = line.strip()
        if not line: continue
        vm = _UNIT_RE.search(line)
        if not vm: continue
        value     = float(vm.group(1)); unit = vm.group(2).strip()
        name_part = re.sub(r"[:\|\-|\*\.]+", "", line[:vm.start()].strip()).strip().lower()
        name_part = re.sub(r"\s+", " ", name_part)
        if (not name_part or len(name_part) < 2 or len(name_part) > 50
                or name_part in _SKIP or re.match(r"^\d", name_part)):
            continue
        labs[name_part] = {"value": value, "unit": unit}
    for name in ner_names:
        if name in labs: continue
        m = re.search(re.escape(name) + r"[^\n]{0,60}", text, re.I)
        if m:
            vm = _UNIT_RE.search(m.group(0))
            if vm: labs[name] = {"value": float(vm.group(1)), "unit": vm.group(2).strip()}
    return labs

def extract_vitals(text):
    vitals = {}
    patterns = [
        (r"(?:bp|blood\s*pressure)[:\s]+(\d{2,3})\s*/\s*(\d{2,3})\s*(?:mmhg)?", "bp"),
        (r"(?:hr|heart\s*rate|pulse)[:\s=]+(\d{2,3})\s*(?:bpm|/min)?",  "heart rate",       "bpm"),
        (r"(?:bmi|body\s*mass\s*index)[:\s=]+([\d.]+)\s*(?:kg/m2)?",    "bmi",              "kg/m2"),
        (r"(?:wt|weight)[:\s=]+([\d.]+)\s*(kg|lbs?)",                   "weight",           "kg"),
        (r"(?:ht|height)[:\s=]+([\d.]+)\s*(cm|m\b|ft|in)",              "height",           "cm"),
        (r"(?:temp|temperature)[:\s=]+([\d.]+)\s*(deg\s*[CF]|°[CF])",   "temperature",      "deg C"),
        (r"(?:spo2|o2\s*sat|oxygen\s*sat)[:\s=]+([\d.]+)\s*%?",         "spo2",             "%"),
        (r"(?:rr|resp(?:iratory)?\s*rate)[:\s=]+(\d{1,2})\s*(?:/min)?", "respiratory rate", "/min"),
    ]
    for item in patterns:
        pat = item[0]; name = item[1]
        m   = re.search(pat, text, re.I)
        if not m: continue
        if name == "bp":
            vitals["blood pressure systolic"]  = {"value": float(m.group(1)), "unit": "mmHg"}
            vitals["blood pressure diastolic"] = {"value": float(m.group(2)), "unit": "mmHg"}
        else:
            default_unit = item[2] if len(item) > 2 else ""
            try:
                unit = (m.group(2).strip() if m.lastindex and m.lastindex >= 2 and m.group(2)
                        else default_unit)
            except Exception:
                unit = default_unit
            vitals[name] = {"value": float(m.group(1)), "unit": unit}
    return vitals

def extract_treatments(text, bert_ents):
    treatments = set()
    for ent in bert_ents:
        lbl = ent.get("entity_group", ent.get("label", "")).upper()
        if any(x in lbl for x in ("DRUG", "CHEM", "CHEMICAL")):
            name = ent["word"].strip().lower()
            if len(name) > 2: treatments.add(name)
    for pat in [
        r"(?:prior|previous|past)\s+(?:treatment|therapy|chemo(?:therapy)?)[s]?[:\s]+([^\n]+)",
        r"(?:received|underwent|history of)[:\s]+([^\n,;]{5,60})",
        r"(?:medications?|current\s+meds?)[:\s]+([^\n]+)",
    ]:
        for m in re.finditer(pat, text, re.I):
            for part in re.split(r"[,;]|\band\b", m.group(1)):
                part = re.sub(
                    r"^(prior|previous|received|underwent|history of|medications?)[:\s]+",
                    "", part.strip().lower(), flags=re.I).strip()
                if 2 < len(part) < 60: treatments.add(part)
    return sorted(treatments)

def extract_comorbidities(text, primary_conditions):
    comorbidities = set()
    for pat in [
        r"(?:pmh|past medical history|medical history|comorbidities?|other conditions?)[:\s]+([^\n]{5,200})",
        r"(?:known case of|background history of)[:\s]+([^\n,;]+)",
    ]:
        m = re.search(pat, text, re.I)
        if m:
            for part in re.split(r"[,;]|\band\b", m.group(1)):
                part = re.sub(r"[\.]+", "", part.strip().lower()).strip()
                if 3 < len(part) < 80 and part not in primary_conditions:
                    comorbidities.add(part)
    return sorted(comorbidities)

def structure_patient(text):
    bert_ents     = _get_bert_ents(text[:8000])
    age, sex      = extract_demographics(text)
    conditions    = extract_conditions(text, bert_ents)
    ecog          = extract_ecog(text)
    lab_tests     = extract_lab_tests(text, bert_ents)
    vitals        = extract_vitals(text)
    treatments    = extract_treatments(text, bert_ents)
    comorbidities = extract_comorbidities(text, set(conditions))
    patient_id    = extract_patient_id(text)
    for vname in vitals: lab_tests.pop(vname, None)
    return {
        "patient_id":       patient_id,
        "age":              age,
        "sex":              sex,
        "conditions":       conditions,
        "ecog":             ecog,
        "lab_tests":        lab_tests,
        "vitals":           vitals,
        "prior_treatments": treatments,
        "comorbidities":    comorbidities,
    }

def validate(patient):
    issues = []
    if not patient.get("age"):        issues.append("age not found")
    if not patient.get("sex"):        issues.append("sex not found")
    if not patient.get("conditions"): issues.append("no conditions found")
    if not patient.get("lab_tests"):  issues.append("no lab tests found")
    return issues


# ---------------------------------------------------------------------------
# HTML result renderer  (ported from notebook Gradio section → Streamlit)
# ---------------------------------------------------------------------------
def _get_styles():
    return """
<style>

@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@300;400;500;600&display=swap');

:root {
    --bg:          #f8fafc;
    --bg-card:     #ffffff;
    --bg-inner:    #ffffff;
    --border:      #e5e7eb;
    --border-soft: #edf1f5;
    --text:        #111827;
    --text-muted:  #374151;
    --text-dim:    #6b7280;
    --text-faint:  #9ca3af;
    --blue:        #2563eb;
    --blue-bg:     #dbeafe;
    --blue-bdr:    #93c5fd;
    --green:       #16a34a;
    --green-bg:    #dcfce7;
    --green-bdr:   #86efac;
    --yellow:      #ca8a04;
    --yellow-bg:   #fef9c3;
    --yellow-bdr:  #fde68a;
    --red:         #dc2626;
    --red-bg:      #fee2e2;
    --red-bdr:     #fca5a5;
    --shadow:      rgba(0,0,0,0.08);
}

@media (prefers-color-scheme: dark) {
    :root {
        --bg:          #0C0C0E;
        --bg-card:     #141416;
        --bg-inner:    #1A1A1E;
        --border:      rgba(255,255,255,0.14);
        --border-soft: rgba(255,255,255,0.08);
        --text:        #F2F0EB;
        --text-muted:  #BCBAB4;
        --text-dim:    #9A9894;
        --text-faint:  #6B6A66;
        --blue:        #6FA8FF;
        --blue-bg:     rgba(111,168,255,0.12);
        --blue-bdr:    rgba(111,168,255,0.35);
        --green:       #3EE09A;
        --green-bg:    rgba(62,224,154,0.12);
        --green-bdr:   rgba(62,224,154,0.35);
        --yellow:      #F5A623;
        --yellow-bg:   rgba(245,166,35,0.12);
        --yellow-bdr:  rgba(245,166,35,0.35);
        --red:         #FF5C5C;
        --red-bg:      rgba(255,92,92,0.12);
        --red-bdr:     rgba(255,92,92,0.35);
        --shadow:      rgba(0,0,0,0.5);
    }
}

html, body {
    margin:0;
    padding:0;
    background:var(--bg);
}

.ctm-root {
    font-family:'IBM Plex Sans',sans-serif;
    background:var(--bg);
    color:var(--text);
    padding:16px;
    margin:0;
    border-radius:16px;
}

.section-heading {
    font-family:'IBM Plex Mono',monospace;
    font-size:11px; font-weight:600; text-transform:uppercase;
    letter-spacing:2px; color:var(--text-dim);
    padding-bottom:12px; border-bottom:1px solid var(--border-soft);
    margin-bottom:18px;
}

.trial-card {
    background:var(--bg-card);
    border:1px solid var(--border);
    border-radius:18px;
    margin-bottom:16px;
    overflow:hidden;
    transition:all 0.18s ease;
}
.trial-card:hover {
    transform:translateY(-2px);
    box-shadow:0 10px 28px var(--shadow);
}

.card-header {
    padding:20px 22px; display:flex; align-items:flex-start;
    gap:18px; cursor:pointer; list-style:none;
}

.rank-badge {
    font-family:'IBM Plex Mono',monospace;
    font-size:34px; font-weight:700; color:var(--text-dim); min-width:44px;
}

.card-main { flex:1; }

.trial-title {
    font-size:17px; font-weight:600; color:var(--text);
    margin:0 0 12px 0; line-height:1.45;
}

.card-meta { display:flex; flex-wrap:wrap; gap:8px; align-items:center; }

.badge {
    font-size:11px; font-family:'IBM Plex Mono',monospace;
    padding:4px 10px; border-radius:999px; font-weight:600;
}

.badge-green  { background:var(--green-bg);  color:var(--green);  border:1px solid var(--green-bdr); }
.badge-blue   { background:var(--blue-bg);   color:var(--blue);   border:1px solid var(--blue-bdr); }
.badge-yellow { background:var(--yellow-bg); color:var(--yellow); border:1px solid var(--yellow-bdr); }
.badge-red    { background:var(--red-bg);    color:var(--red);    border:1px solid var(--red-bdr); }
.badge-gray   { background:var(--bg-inner);  color:var(--text-dim); border:1px solid var(--border); }

.trial-id { font-family:'IBM Plex Mono',monospace; font-size:11px; color:var(--text-faint); }

.score-arc { display:flex; align-items:center; gap:8px; margin-left:auto; }

.score-circle {
    width:62px; height:62px; border-radius:50%;
    display:flex; align-items:center; justify-content:center;
    font-family:'IBM Plex Mono',monospace; font-size:14px; font-weight:700;
}
.score-hi  { background:var(--green-bg);  border:2px solid var(--green);  color:var(--green); }
.score-mid { background:var(--yellow-bg); border:2px solid var(--yellow); color:var(--yellow); }
.score-lo  { background:var(--red-bg);    border:2px solid var(--red);    color:var(--red); }

.card-detail { border-top:1px solid var(--border-soft); padding:20px 22px; }

.summary-box {
    background:var(--bg-inner);
    border-left:4px solid var(--blue);
    border-radius:0 10px 10px 0;
    padding:14px 16px; font-size:14px; color:var(--text-muted);
    line-height:1.7; margin-bottom:18px;
}

.criteria-grid {
    display:grid; grid-template-columns:repeat(auto-fill,minmax(260px,1fr));
    gap:12px; margin-bottom:18px;
}

.crit-row {
    background:var(--bg-inner); border:1px solid var(--border-soft);
    border-radius:10px; padding:12px 14px;
}
.crit-pass    { border-left:4px solid var(--green); }
.crit-partial { border-left:4px solid var(--yellow); }
.crit-fail    { border-left:4px solid var(--red); }
.crit-unknown { border-left:4px solid var(--text-dim); }

.crit-name {
    font-size:11px; font-weight:700; color:var(--text);
    text-transform:uppercase; letter-spacing:0.6px; margin-bottom:6px;
}
.crit-values { display:flex; gap:8px; font-size:13px; flex-wrap:wrap; }
.crit-patient { color:var(--text-muted); }
.crit-trial   { color:var(--text-faint); }

.verdict-box {
    display:inline-flex; align-items:center; gap:8px;
    padding:10px 16px; border-radius:999px;
    font-family:'IBM Plex Mono',monospace; font-weight:700; font-size:12px;
    margin-top:12px;
}
.verdict-eligible   { background:var(--green-bg); border:1px solid var(--green); color:var(--green); }
.verdict-ineligible { background:var(--red-bg);   border:1px solid var(--red);   color:var(--red); }

.no-match {
    background:var(--red-bg); border:1px solid var(--red-bdr);
    border-radius:16px; padding:28px; text-align:center;
    color:var(--red); font-size:15px;
}

.disclaimer {
    margin-top:20px; background:var(--bg-card); border:1px solid var(--border);
    border-radius:12px; padding:14px 16px; font-size:11px;
    color:var(--text-faint); line-height:1.6;
}

details > summary { list-style:none; }
details > summary::-webkit-details-marker { display:none; }

.chevron {
    display:inline-block; transition:transform 0.2s;
    color:var(--text-dim); margin-left:8px; font-size:12px;
}
details[open] .chevron { transform:rotate(90deg); }

</style>
"""

def _status_badge(label):
    label = label or ""
    cls = ("badge-green"  if "Recruiting" in label and "Not" not in label else
           "badge-blue"   if "Opening" in label or "Invitation" in label else
           "badge-yellow" if "Active" in label else
           "badge-red"    if "Completed" in label or "Terminated" in label else
           "badge-gray")
    return f'<span class="badge {cls}">{label}</span>'

def _score_circle(score):
    try: s = float(str(score).replace("%", "").strip())
    except Exception: s = 0
    cls = "score-hi" if s >= 75 else "score-mid" if s >= 50 else "score-lo"
    return f'<div class="score-circle {cls}">{score}%</div>'

def _rank_badge(rank):
    return f'<div class="rank-badge">#{rank}</div>'

def _crit_row_html(name, detail, status_class):
    pval  = detail.get("patient", detail.get("patient_value", "—"))
    tval  = detail.get("trial",   detail.get("trial_req", "—"))
    overlap = detail.get("overlap", "")
    extra = f' <span style="color:var(--text-faint);font-size:11px;">({overlap})</span>' if overlap else ""
    return f"""<div class="crit-row crit-{status_class}">
      <div>
        <div class="crit-name">{name}</div>
        <div class="crit-values">
          <span class="crit-patient">{pval}{extra}</span>
          <span class="crit-arrow">→</span>
          <span class="crit-trial">trial: {tval}</span>
        </div>
      </div>
    </div>"""

def _render_dt(rules, verdict, confidence):
    vcls = "verdict-eligible" if "ELIGIBLE" in str(verdict).upper() and "NOT" not in str(verdict).upper() else "verdict-ineligible"
    icon = "✓" if "eligible" in vcls else "✗"
    return f'<div class="verdict-box {vcls}">{icon} {verdict}</div>'

def _render_match(m):
    if not m or not isinstance(m, dict): return ""
    exp    = m.get("explanation") or {}
    rank   = m.get("rank", 1)
    score  = m.get("score", "?")
    title  = m.get("title", "Unknown Trial")
    tid    = m.get("trial_id", "")
    slabel = m.get("status_label", m.get("status", ""))
    label  = m.get("label", "")
    xgb_v  = m.get("xgb_verdict", "")

    summary  = exp.get("summary", "")
    criteria = exp.get("criteria_breakdown", {})
    dt_verd  = exp.get("dt_verdict", "")

    # Split criteria into groups
    passes   = {k: v for k, v in criteria.items() if v.get("status") == "PASS"}
    partials = {k: v for k, v in criteria.items() if v.get("status") == "PARTIAL"}
    fails    = {k: v for k, v in criteria.items() if v.get("status") == "FAIL"}
    unknowns = {k: v for k, v in criteria.items() if v.get("status") == "UNKNOWN"}

    def _section(label_text, status_class, items):
        if not items:
            return ""
        rows = "".join(_crit_row_html(k, v, status_class) for k, v in items.items())
        color = {"pass": "var(--green)", "partial": "var(--yellow)",
                 "fail": "var(--red)", "unknown": "var(--text-dim)"}[status_class]
        return f"""
            <div style="font-size:10px;font-weight:600;letter-spacing:1.5px;
                        text-transform:uppercase;margin:14px 0 8px 0;
                        color:{color};">{label_text}</div>
            <div class="criteria-grid">{rows}</div>"""

    crit_html = (
        _section("Matched", "pass", passes) +
        _section("Partial Match", "partial", partials) +
        _section("Not Met", "fail", fails) +
        _section("Unknown", "unknown", unknowns)
    )

    xgb_cls  = "xgb-eligible" if xgb_v and "ELIGIBLE" in xgb_v and "NOT" not in xgb_v else "xgb-ineligible"
    xgb_html = f'<span class="xgb-label {xgb_cls}">{xgb_v}</span>' if xgb_v else ""

    if dt_verd:
        vcls = "verdict-eligible" if "ELIGIBLE" in dt_verd.upper() and "NOT" not in dt_verd.upper() else "verdict-ineligible"
        icon = "✓" if vcls == "verdict-eligible" else "✗"
        verdict_html = f'<div class="verdict-box {vcls}">{icon} {dt_verd}</div>'
    else:
        verdict_html = ""

    detail_html = ""
    if summary or crit_html or verdict_html:
        detail_html = f"""<div class="card-detail">
          {'<div class="summary-box">' + summary + '</div>' if summary else ''}
          {'<div class="section-heading">Criteria Breakdown</div>' + crit_html if crit_html else ''}
          {verdict_html}
        </div>"""

    return f"""<details class="trial-card">
      <summary class="card-header">
        {_rank_badge(rank)}
        <div class="card-main">
          <div class="trial-title">{title}</div>
          <div class="card-meta">
            {_status_badge(slabel)}
            <span class="badge badge-gray">{label}</span>
            {xgb_html}
            <span class="trial-id">{tid}</span>
          </div>
        </div>
        <div class="score-arc">{_score_circle(score)}<span class="chevron">▶</span></div>
      </summary>
      {detail_html}
    </details>"""


def render_results(result):
    """Convert match_patient() result → rich HTML for st.components.v1.html()."""
    if not result:
        return "<div class='ctm-root'><div class='no-match'>No result to display.</div></div>"

    patient_id = result.get("patient_id", "Unknown")
    age        = result.get("age", "?")
    sex        = result.get("sex", "?")
    conditions = result.get("conditions", [])
    ecog       = result.get("ecog", "?")
    matches    = [m for m in result.get("top_matches", []) if isinstance(m, dict)]
    no_match   = result.get("no_match", False)
    disclaimer = result.get("disclaimer", "")



    if no_match or not matches:
        reason    = result.get("no_match_reason", "No suitable trials found.")
        body_html = f'<div class="no-match"><div style="font-size:36px;margin-bottom:12px">🔴</div><strong>No Matching Trial Found</strong><br><br>{reason}</div>'
    else:
        heading   = '<div class="section-heading">Top Matched Trials — click any card to expand details</div>'
        body_html = heading + "".join(_render_match(m) for m in matches)

    disc_html = f'<div class="disclaimer">⚠ {disclaimer}</div>' if disclaimer else ""
    return f'<div class="ctm-root">{_get_styles()}{body_html}{disc_html}</div>'


# ---------------------------------------------------------------------------
# Main pipeline entry point
# ---------------------------------------------------------------------------
def run_pipeline(uploaded_file):
    import tempfile

    _ensure_models_loaded()
    trials = load_trials("trials_extracted.json")

    # ── Manual entry: uploaded_file is already a structured dict ──
    if isinstance(uploaded_file, dict):
        manual = uploaded_file

        # Normalize gender → sex
        sex = manual.get("Gender", manual.get("sex", "")).upper()
        if sex == "MALE":   sex = "MALE"
        elif sex == "FEMALE": sex = "FEMALE"
        else: sex = None

        # Build conditions list from diagnosis string
        diagnosis = manual.get("Primary Diagnosis", "")
        conditions = [diagnosis.lower().strip()] if diagnosis else []

        # Pull labs — already flattened into top-level keys by app.py
        lab_keys = {"Hemoglobin","Platelets","WBC","Creatinine","eGFR",
                    "AST","ALT","Albumin","PSA","LDH","HbA1c","TSH","BNP"}
        lab_tests = {}
        for k, v in manual.items():
            if k in lab_keys or (isinstance(manual.get("Lab Values"), dict) and k in manual["Lab Values"]):
                try:
                    lab_tests[k.lower()] = {"value": float(v), "unit": ""}
                except (ValueError, TypeError):
                    pass

        # Vitals
        vitals = {}
        for src_key, dest_key in [
            ("Heart Rate", "heart rate"),
            ("SpO2",       "spo2"),
            ("Temperature","temperature"),
            ("Weight",     "weight"),
            ("Height",     "height"),
        ]:
            if manual.get(src_key):
                try:
                    vitals[dest_key] = {"value": float(manual[src_key]), "unit": ""}
                except (ValueError, TypeError):
                    pass
        if manual.get("Blood Pressure"):
            try:
                sys, dia = manual["Blood Pressure"].split("/")
                vitals["blood pressure systolic"]  = {"value": float(sys), "unit": "mmHg"}
                vitals["blood pressure diastolic"] = {"value": float(dia), "unit": "mmHg"}
            except Exception:
                pass

        # Prior treatments
        raw_tx = manual.get("Prior Treatments", "")
        prior_treatments = [t.strip().lower() for t in re.split(r"[,;]", raw_tx) if t.strip()] if raw_tx else []

        # ECOG
        try:
            ecog = int(manual.get("ECOG", 0))
        except (ValueError, TypeError):
            ecog = None

        patient = {
            "patient_id":       "MANUAL_ENTRY",
            "age":              manual.get("Age"),
            "sex":              sex,
            "conditions":       conditions,
            "ecog":             ecog,
            "lab_tests":        lab_tests,
            "vitals":           vitals,
            "prior_treatments": prior_treatments,
            "comorbidities":    [],
        }

        result = match_patient(
            patient, trials, XGB, DT, SCALER_FULL, ALL_FEAT_NAMES, top_n=10)
        return patient, result

    # ── File upload: original PDF path ──
    suffix = Path(uploaded_file.name).suffix or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        temp_path = tmp.name

    try:
        raw_text   = extract_text(temp_path)
        clean_text = preprocess(raw_text)
        patient    = structure_patient(clean_text)
        result     = match_patient(
            patient, trials, XGB, DT, SCALER_FULL, ALL_FEAT_NAMES, top_n=10)
    finally:
        os.unlink(temp_path)

    return patient, result